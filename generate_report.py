import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_run_font(run, font_name='Times New Roman', size=13):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0) # Black

def add_paragraph(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run)
    return p

def main():
    doc = Document()
    
    # Configure default style to Times New Roman, 13pt, Black
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run('BÁO CÁO CẬP NHẬT CHỨC NĂNG HỆ THỐNG')
    title_run.bold = True
    set_run_font(title_run, size=16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub_run = sub.add_run('Chuyên đề: Hoàn thiện tính năng Đồng bộ Dữ liệu Tự động và Giao diện Quản trị')
    sub_run.bold = True
    set_run_font(sub_run, size=14)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # spacing
    
    # Section 1
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run('1. CÁC TÍNH NĂNG ĐƯỢC BỔ SUNG MỚI')
    r1.bold = True
    set_run_font(r1, size=14)
    
    add_paragraph(doc, '- Cơ sở dữ liệu: Bổ sung thêm trường trạng thái "Đang đồng bộ" vào bảng Cấu hình hệ thống nhằm theo dõi chính xác tiến trình chạy ngầm.', style='List Paragraph')
    add_paragraph(doc, '- Xử lý đa luồng (Chạy ngầm): Xây dựng hệ thống tự động quét và lập lịch định kỳ để đồng bộ dữ liệu mà không làm treo hay đơ trang web của người dùng.', style='List Paragraph')
    add_paragraph(doc, '- Cơ chế làm mới giao diện tự động: Trang danh sách cấu hình sẽ tự động cập nhật số liệu 10 giây một lần. Quản trị viên không cần phải tải lại trang thủ công mà vẫn thấy được số lần đồng bộ thay đổi liên tục.', style='List Paragraph')
    
    # Section 2
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run('2. NHỮNG ĐIỀU CHỈNH VÀ NÂNG CẤP GIAO DIỆN')
    r2.bold = True
    set_run_font(r2, size=14)
    
    add_paragraph(doc, '- Nâng cấp Biểu mẫu thiết lập thời gian:', style='List Paragraph')
    add_paragraph(doc, '+ Chuyển đổi hộp kiểm (checkbox) thông thường thành nút gạt (toggle) trực quan hơn.', style='List Bullet 2')
    add_paragraph(doc, '+ Thiết kế lại khu vực nhập chu kỳ đồng bộ. Dù tính năng đang tắt, biểu mẫu vẫn hiển thị rõ ràng thông số (có làm mờ nhẹ để phân biệt) nhằm giúp quản trị viên dễ quan sát.', style='List Bullet 2')
    add_paragraph(doc, '+ Thay vì phải tự quy đổi thời gian ra phút, hệ thống đã cho phép chọn đơn vị trực tiếp như: Phút, Giờ, Ngày, Tháng, Năm.', style='List Bullet 2')
    
    add_paragraph(doc, '- Cải thiện trạng thái hiển thị quá trình đồng bộ:', style='List Paragraph')
    add_paragraph(doc, '+ Ngay khi hệ thống tự động chạy, nhãn hiển thị sẽ đổi sang màu cam với dòng chữ "Đang đồng bộ...".', style='List Bullet 2')
    add_paragraph(doc, '+ Nút bấm "Đồng bộ" thủ công sẽ bị khoá tạm thời trong suốt quá trình này. Việc này nhằm ngăn chặn thao tác bấm liên tục nhiều lần gây quá tải máy chủ và kẹt dữ liệu.', style='List Bullet 2')
    
    # Section 3
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run('3. CÁC THÀNH PHẦN ĐÃ ĐƯỢC LƯỢC BỎ')
    r3.bold = True
    set_run_font(r3, size=14)
    
    add_paragraph(doc, '- Xoá bỏ các mã lệnh giao diện cũ liên quan đến việc ẩn/hiện cứng nhắc của ô nhập thời gian.', style='List Paragraph')
    add_paragraph(doc, '- Gỡ bỏ các định dạng làm mờ ô chữ màu xám khiến quản trị viên khó đọc thông tin khi tính năng đang tắt.', style='List Paragraph')
    
    # Section 4
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run('4. QUY TRÌNH XỬ LÝ LÔ-GÍC CỦA HỆ THỐNG')
    r4.bold = True
    set_run_font(r4, size=14)
    
    add_paragraph(doc, 'Hệ thống áp dụng phương thức Cập nhật Thông minh (Cập nhật hoặc Tạo mới) để kiểm soát dữ liệu:', style='List Paragraph')
    add_paragraph(doc, '+ Đối với dữ liệu đã tồn tại: Căn cứ vào mã định danh, nếu phát hiện có sự khác biệt, hệ thống sẽ ghi đè các thông số cũ bằng thông tin mới nhất từ máy chủ gốc.', style='List Bullet 2')
    add_paragraph(doc, '+ Đối với dữ liệu mới: Hệ thống sẽ tự động khởi tạo và bổ sung vào danh mục.', style='List Bullet 2')
    add_paragraph(doc, '+ Đối với dữ liệu bị gỡ bỏ từ máy chủ gốc: Hệ thống sẽ không xoá thiết bị này. Việc giữ lại dữ liệu nhằm bảo vệ tính toàn vẹn của các chứng từ, hồ sơ cũ đang liên kết với thiết bị đó.', style='List Bullet 2')
    
    add_paragraph(doc, 'Xử lý kết nối mạng và bảo mật:', style='List Paragraph')
    add_paragraph(doc, '+ Hệ thống tự động chuyển đổi sang kết nối mã hoá bảo mật (chặn lỗi tường lửa).', style='List Bullet 2')
    add_paragraph(doc, '+ Tích hợp cơ chế tự động nhận diện và mã hoá tài khoản đăng nhập khi cung cấp thông tin tiêu đề kết nối (Header).', style='List Bullet 2')
    
    # Section 5: Database Diagram
    h5 = doc.add_heading(level=1)
    r5 = h5.add_run('5. SƠ ĐỒ CẤU TRÚC CƠ SỞ DỮ LIỆU ĐÃ THAY ĐỔI')
    r5.bold = True
    set_run_font(r5, size=14)
    
    add_paragraph(doc, 'Bảng cấu trúc: Bảng Cấu hình hệ thống', bold=True)
    
    # Create Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_titles = ['Tên Trường', 'Kiểu Dữ Liệu', 'Ràng buộc', 'Ghi chú']
    for i, title in enumerate(hdr_titles):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        set_run_font(run)
    
    # Data Rows
    db_schema = [
        ('key', 'Văn bản (255)', 'Bắt buộc, Duy nhất', 'Mã cấu hình (VD: API_TRAM)'),
        ('value', 'Văn bản', 'Bắt buộc', 'Đường dẫn liên kết dữ liệu'),
        ('description', 'Văn bản', 'Tuỳ chọn', 'Mô tả tính năng'),
        ('auto_sync_enabled', 'Đúng/Sai (Boolean)', 'Mặc định: Sai', 'Trạng thái Bật/Tắt chạy tự động'),
        ('sync_interval_minutes', 'Số nguyên', 'Mặc định: 1440', 'Chu kỳ lặp lại (Tính bằng Phút)'),
        ('last_sync_time', 'Ngày giờ', 'Tuỳ chọn', 'Thời điểm chạy thành công lần cuối'),
        ('last_sync_status', 'Văn bản (50)', 'Tuỳ chọn', 'Trạng thái lần cuối (Thành công/Lỗi)'),
        ('sync_count', 'Số nguyên', 'Mặc định: 0', 'Tổng số lần đã chạy'),
        ('is_syncing', 'Đúng/Sai (Boolean)', 'Mặc định: Sai', '[Bổ sung mới] Trạng thái đang tải ngầm'),
    ]
    
    for row_data in db_schema:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run)

    file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Bao_Cao_Hoan_Thien_He_Thong.docx')
    doc.save(file_path)
    print(f"Created report at {file_path}")

if __name__ == '__main__':
    main()
