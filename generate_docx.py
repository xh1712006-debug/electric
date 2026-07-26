import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font_black(run):
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_heading_black(heading):
    for run in heading.runs:
        set_font_black(run)

def create_document():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('TÀI LIỆU PHÂN TÍCH VAI TRÒ VÀ LUỒNG HOẠT ĐỘNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_black(title)
    
    p = doc.add_paragraph('Dự án: Hệ thống Quản lý Rơ-le Kỹ thuật số')
    for run in p.runs: set_font_black(run)
    
    p = doc.add_paragraph('--------------------------------------------------')
    for run in p.runs: set_font_black(run)
    
    # Section 1: Roles
    h1 = doc.add_heading('1. CÁC VAI TRÒ VÀ CHỨC NĂNG', level=1)
    set_heading_black(h1)
    
    roles = [
        {
            'name': 'Quản trị viên',
            'permissions': ['Xem danh sách trạm', 'Quản lý người dùng', 'Tạo phiếu cấu hình', 'Phê duyệt phiếu', 'Phân công phiếu', 'Thực thi cấu hình', 'Giám sát cấu hình'],
            'desc': 'Toàn quyền kiểm soát hệ thống. Quản lý tài khoản, phân quyền, cấu hình hệ thống và có thể thực hiện mọi thao tác của các vai trò khác (Duyệt phiếu, rà soát, cấu hình).'
        },
        {
            'name': 'Điều phối viên',
            'permissions': ['Xem danh sách trạm', 'Tạo phiếu cấu hình', 'Phân công phiếu'],
            'desc': 'Tạo phiếu cấu hình mới, sử dụng công nghệ nhận dạng quang học để trích xuất số liệu, rà soát tính hợp lệ của phiếu và phân công công việc về cho các trạm hoặc Kỹ thuật viên.'
        },
        {
            'name': 'Trưởng nhóm Trạm',
            'permissions': ['Xem danh sách trạm', 'Phân công phiếu'],
            'desc': 'Quản lý chung tại trạm. Có trách nhiệm tiếp nhận phân công từ Điều phối viên và điều phối, giao việc lại cho Kỹ thuật viên trong trạm của mình.'
        },
        {
            'name': 'Kỹ thuật viên',
            'permissions': ['Thực thi cấu hình'],
            'desc': 'Trực tiếp xuống trạm thực hiện việc cấu hình rơ-le. Sau khi hoàn tất công việc, Kỹ thuật viên sẽ nhập thông số thực tế và thực hiện ký số để xác nhận hoàn thành.'
        },
        {
            'name': 'Giám sát trạm',
            'permissions': ['Giám sát cấu hình'],
            'desc': 'Giám sát công việc tại trạm. Có trách nhiệm rà soát các thông số mà Kỹ thuật viên vừa nhập, kiểm tra tính chính xác và ký số xác nhận nghiệm thu.'
        }
    ]
    
    for r in roles:
        p = doc.add_paragraph()
        runner = p.add_run(r['name'])
        runner.bold = True
        runner.font.size = Pt(12)
        set_font_black(runner)
        
        p_desc = doc.add_paragraph(f"Chức năng/Mô tả: {r['desc']}", style='List Bullet')
        for run in p_desc.runs: set_font_black(run)
            
        p_perm = doc.add_paragraph(f"Quyền hạn: {', '.join(r['permissions'])}", style='List Bullet')
        for run in p_perm.runs: set_font_black(run)
    
    # Section 2: Main Workflows
    h2 = doc.add_heading('2. CÁC LUỒNG HOẠT ĐỘNG CHÍNH', level=1)
    set_heading_black(h2)
    
    h2_1 = doc.add_heading('2.1. Quy trình Phiếu Cấu Hình Rơ-le', level=2)
    set_heading_black(h2_1)
    
    p = doc.add_paragraph('Đây là luồng nghiệp vụ chính để cấu hình một rơ-le, trải qua quá trình ký số đa bước:')
    for run in p.runs: set_font_black(run)
    
    steps = [
        'Bản nháp: Điều phối viên tạo phiếu và quét tài liệu thông số.',
        'Chờ rà soát: Phiếu được tạo và chờ kiểm tra tính hợp lệ.',
        'Đã chuyển về Trạm: Điều phối viên gửi phiếu về trạm cụ thể.',
        'Đã giao cho Kỹ thuật viên: Trưởng nhóm trạm phân công cho Kỹ thuật viên.',
        'Đang thực hiện: Kỹ thuật viên tiếp nhận và tiến hành cấu hình rơ-le thực tế.',
        'Chờ duyệt ban hành: Kỹ thuật viên hoàn tất, Giám sát viên kiểm tra và ký số xác nhận. Sau đó chờ Quản trị viên hoặc cấp trên phê duyệt.',
        'Hoàn thành: Phiếu được duyệt và đóng, cấu hình chính thức có hiệu lực.'
    ]
    
    for step in steps:
        p_step = doc.add_paragraph(step, style='List Number')
        for run in p_step.runs: set_font_black(run)
        
    p = doc.add_paragraph('Quá trình ký số được thực hiện nghiêm ngặt ở các bước thực thi (Kỹ thuật viên), nghiệm thu (Giám sát) và phê duyệt (Quản trị viên).')
    for run in p.runs: set_font_black(run)
    
    h2_2 = doc.add_heading('2.2. Quy trình Tự động Kiểm tra Rơ-le', level=2)
    set_heading_black(h2_2)
    
    p = doc.add_paragraph('Hệ thống có cơ chế theo dõi tự động định kỳ các thông số của rơ-le để so sánh với tiêu chuẩn:')
    for run in p.runs: set_font_black(run)
    
    auto_steps = [
        'Hệ thống định kỳ lấy dữ liệu thực tế từ rơ-le thông qua giao thức mạng.',
        'So sánh dữ liệu thực tế với thông số chuẩn đã lưu trong hệ thống.',
        'Ghi nhận lại trạng thái: Trùng khớp, Có sai lệch, hoặc Lỗi kết nối.',
        'Nếu có sai lệch, hệ thống tự động tạo một Phiếu xử lý sự cố.'
    ]
    for step in auto_steps:
        p_step = doc.add_paragraph(step, style='List Bullet')
        for run in p_step.runs: set_font_black(run)
        
    h2_3 = doc.add_heading('2.3. Quy trình Xử lý Sự cố Rơ-le', level=2)
    set_heading_black(h2_3)
    
    p = doc.add_paragraph('Khi có Phiếu xử lý sự cố được tạo ra do sai lệch thông số, quy trình xử lý trải qua các bước:')
    for run in p.runs: set_font_black(run)
    
    ticket_steps = [
        'Phân phối viên tiếp nhận và xử lý sự cố',
        'Điều chuyển về Trạm xử lý',
        'Kỹ thuật viên trực tiếp khắc phục',
        'Giám sát kiểm tra và ký xác nhận',
        'Quản trị viên ký duyệt hoàn thành',
        'Hoàn tất và đóng phiếu sự cố'
    ]
    for step in ticket_steps:
        p_step = doc.add_paragraph(step, style='List Number')
        for run in p_step.runs: set_font_black(run)

    p = doc.add_paragraph('\n-- Hết --', style='Normal')
    for run in p.runs: set_font_black(run)
    
    # Xóa style màu xanh mặc định của word cho tất cả các heading
    for style in doc.styles:
        if style.name.startswith('Heading'):
            if style.font.color:
                style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Save document
    doc.save('Vai_Tro_Va_Quy_Trinh_EVN_RMS.docx')

if __name__ == "__main__":
    create_document()
