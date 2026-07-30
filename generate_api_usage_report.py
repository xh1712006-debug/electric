import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def set_run_font(run, font_name='Times New Roman', size=13):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0) # Black

def add_paragraph(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run)
    return p

def main():
    doc = Document()
    
    # Configure default style to Times New Roman, 13pt, Black
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run('BÁO CÁO ÁP DỤNG DỮ LIỆU TÍCH HỢP TỪ HỆ THỐNG PMIS (API)')
    title_run.bold = True
    set_run_font(title_run, size=16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub_run = sub.add_run('Mục đích: Tổng hợp mục đích sử dụng các dữ liệu được đồng bộ từ API vào các chức năng nghiệp vụ của dự án.')
    sub_run.bold = True
    set_run_font(sub_run, size=14)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # spacing
    
    # Section 1
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run('1. TỔNG QUAN CÁC NHÓM DỮ LIỆU ĐƯỢC ĐỒNG BỘ')
    r1.bold = True
    set_run_font(r1, size=14)
    
    add_paragraph(doc, 'Hệ thống thực hiện kết nối và đồng bộ tự động 11 danh mục dữ liệu cốt lõi từ phần mềm gốc (PMIS). Các danh mục này bao gồm:')
    add_paragraph(doc, '- Nhóm Tổ chức: Đơn vị quản lý, Sở hữu tài sản.', style='List Bullet')
    add_paragraph(doc, '- Nhóm Phân loại kỹ thuật: Cấp điện áp, Loại thiết bị, Trạng thái vận hành.', style='List Bullet')
    add_paragraph(doc, '- Nhóm Nguồn gốc thiết bị: Hãng sản xuất, Nước sản xuất.', style='List Bullet')
    add_paragraph(doc, '- Nhóm Quản lý tài sản lưới điện: Trạm biến áp, Ngăn lộ, Rơ-le, Đường dây, Vị trí, Công trình.', style='List Bullet')
    
    # Section 2
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run('2. ỨNG DỤNG VÀO CÁC CHỨC NĂNG NGHIỆP VỤ')
    r2.bold = True
    set_run_font(r2, size=14)
    
    # 2.1
    p21 = doc.add_paragraph()
    r21 = p21.add_run('2.1. Quản lý Tài sản Lưới điện (Phân hệ Danh mục)')
    r21.bold = True
    set_run_font(r21)
    add_paragraph(doc, 'Dữ liệu API là nền tảng để xây dựng Phân hệ Tài sản lưới điện:', style='List Paragraph')
    add_paragraph(doc, '+ Quản lý Phân cấp Cây thiết bị: Dữ liệu được tổ chức theo mối quan hệ cha-con chặt chẽ (Đơn vị quản lý -> Trạm Biến Áp -> Ngăn lộ -> Rơ-le).', style='List Bullet 2')
    add_paragraph(doc, '+ Tra cứu và Hiển thị: Giúp người dùng có thể tra cứu toàn bộ hồ sơ lý lịch của thiết bị một cách chính xác nhất như năm vận hành, mã định danh, hãng sản xuất.', style='List Bullet 2')
    add_paragraph(doc, '+ Chuẩn hoá Danh pháp: Đảm bảo tên gọi các Trạm, Ngăn lộ, Thiết bị trên hệ thống hoàn toàn khớp với danh pháp kỹ thuật hiện hành của PMIS.', style='List Bullet 2')

    # 2.2
    p22 = doc.add_paragraph()
    r22 = p22.add_run('2.2. Lập Phiếu Chỉnh định Rơ-le (Chức năng Cốt lõi)')
    r22.bold = True
    set_run_font(r22)
    add_paragraph(doc, 'Đây là nghiệp vụ quan trọng nhất của phần mềm. Dữ liệu từ API được dùng trực tiếp để hỗ trợ tạo phiếu:', style='List Paragraph')
    add_paragraph(doc, '+ Tự động Điền thông tin (Auto-fill): Khi kỹ sư chọn một Rơ-le cụ thể để làm Phiếu Chỉnh định, hệ thống sẽ lấy dữ liệu API để tự động điền các thông tin liên quan vào Phiếu (như: Đơn vị quản lý Rơ-le, Thuộc Trạm nào, Ngăn lộ nào, Thuộc cấp điện áp bao nhiêu).', style='List Bullet 2')
    add_paragraph(doc, '+ Liên kết Dữ liệu Vận hành: Gắn kết trạng thái (Đang vận hành/Dự phòng) vào Phiếu để kỹ sư đưa ra quyết định tính toán phù hợp.', style='List Bullet 2')
    add_paragraph(doc, '+ Hạn chế Sai sót Cập nhật thủ công: Tránh được rủi ro kỹ sư gõ nhầm tên Trạm, tên Thiết bị hoặc chọn sai Hãng sản xuất Rơ-le, nhờ đó Form chỉnh định luôn bảo đảm độ chính xác tuyệt đối.', style='List Bullet 2')

    # 2.3
    p23 = doc.add_paragraph()
    r23 = p23.add_run('2.3. Phân quyền và Lọc Dữ liệu (RBAC)')
    r23.bold = True
    set_run_font(r23)
    add_paragraph(doc, 'Dữ liệu Đơn vị Quản lý lấy từ API được sử dụng để làm cơ sở cho ma trận phân quyền:', style='List Paragraph')
    add_paragraph(doc, '+ Kiểm soát Truy cập: Kỹ sư thuộc đơn vị nào (Điện lực Tỉnh/Thành phố) thì chỉ được phép xem, tạo mới và phê duyệt Phiếu Chỉnh định của các thiết bị thuộc đơn vị đó.', style='List Bullet 2')
    add_paragraph(doc, '+ Bộ lọc Thông minh: Dựa trên các thông số API như Cấp điện áp, Trạng thái vận hành để người dùng dễ dàng thu hẹp phạm vi tìm kiếm Phiếu chỉnh định và Thiết bị.', style='List Bullet 2')

    # 2.4
    p24 = doc.add_paragraph()
    r24 = p24.add_run('2.4. Báo cáo Thống kê và Biểu đồ (Dashboard)')
    r24.bold = True
    set_run_font(r24)
    add_paragraph(doc, 'Số liệu tổng quan hiển thị trên Bảng điều khiển (Dashboard) cho Quản lý cấp cao:', style='List Paragraph')
    add_paragraph(doc, '+ Thống kê Thiết bị: Thống kê số lượng Rơ-le theo từng Hãng sản xuất (Siemens, ABB, SEL...), theo Cấp điện áp (110kV, 220kV...).', style='List Bullet 2')
    add_paragraph(doc, '+ Theo dõi Phân mảnh Thiết bị: Nhờ dữ liệu Trạng thái vận hành từ API, Ban Giám đốc dễ dàng nắm bắt có bao nhiêu thiết bị đang bị Hỏng, Đang bảo dưỡng hay Đang Vận hành.', style='List Bullet 2')
    
    # Section 3
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run('3. KẾT LUẬN GIÁ TRỊ MANG LẠI')
    r3.bold = True
    set_run_font(r3, size=14)
    add_paragraph(doc, 'Việc đồng bộ và áp dụng dữ liệu API vào hệ thống mang lại 3 giá trị cốt lõi:')
    add_paragraph(doc, '1. Tính toàn vẹn dữ liệu: Không có sự sai lệch thông tin Thiết bị giữa Hệ thống Quản lý Rơ-le (RMS) và Hệ thống Quản lý Kỹ thuật (PMIS) Quốc gia.')
    add_paragraph(doc, '2. Giảm thiểu thao tác: Người dùng được giải phóng khỏi việc nhập liệu tay các thông tin thiết bị đã có sẵn, đẩy nhanh tiến độ lập Phiếu Chỉnh định.')
    add_paragraph(doc, '3. Quản trị tập trung: Đảm bảo phân quyền bảo mật cấp Đơn vị và Hỗ trợ báo cáo thống kê chính xác tuyệt đối.')

    file_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Bao_Cao_Su_Dung_Du_Lieu_API.docx')
    doc.save(file_path)
    print(f"Created report at {file_path}")

if __name__ == '__main__':
    main()
